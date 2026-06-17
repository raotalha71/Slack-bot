"""
DOCX Generator: Converts Markdown proposal text into a styled Word document.

Features:
  - Cover page with company name and date
  - All proposal sections with proper heading hierarchy
  - Bullet list support
  - Professional styling (Calibri, blue palette)
  - Footer with page numbers
  - Returns bytes — no file written to disk
"""

from __future__ import annotations

import io
import logging
import re
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Color palette (professional blue)
# ---------------------------------------------------------------------------

COLOR_TITLE = RGBColor(0x1A, 0x3A, 0x5C)   # Dark navy
COLOR_H1 = RGBColor(0x1F, 0x5F, 0x8A)      # Medium blue
COLOR_H2 = RGBColor(0x2E, 0x86, 0xAB)      # Light blue
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)    # Dark gray
COLOR_ACCENT = RGBColor(0x1F, 0x5F, 0x8A)  # Medium blue (for bullets)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_font(run, size: int, bold: bool = False, color: RGBColor | None = None):
    """Apply font settings to a run."""
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_paragraph(
    doc: Document,
    text: str,
    style: str = "Normal",
    size: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Add a styled paragraph to the document."""
    para = doc.add_paragraph(style=style)
    para.alignment = alignment
    run = para.add_run(text)
    _set_font(run, size, bold, color or COLOR_BODY)


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading with custom styling."""
    if level == 1:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        _set_font(run, size=18, bold=True, color=COLOR_H1)
        # Add spacing
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_font(run, size=14, bold=True, color=COLOR_H2)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
    else:
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_font(run, size=12, bold=True, color=COLOR_BODY)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)


def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet list item."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    _set_font(run, size=11, color=COLOR_BODY)


def _add_footer(doc: Document) -> None:
    """Add a footer with page numbers."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add separator
    run2 = para.add_run(" | Confidential Proposal")
    _set_font(run2, size=9, color=RGBColor(0x88, 0x88, 0x88))


# ---------------------------------------------------------------------------
# Cover Page
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, client_info: dict[str, Any]) -> None:
    """Add a professional cover page."""
    company = client_info.get("company_name") or "Client"
    today = date.today().strftime("%B %d, %Y")

    # Top spacing
    doc.add_paragraph().paragraph_format.space_before = Pt(60)

    # Main title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("BUSINESS PROPOSAL")
    _set_font(run, size=32, bold=True, color=COLOR_TITLE)

    # Subtitle
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Prepared for {company}")
    _set_font(run, size=20, bold=False, color=COLOR_H1)
    para.paragraph_format.space_before = Pt(12)

    # Divider line
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("─" * 40)
    _set_font(run, size=14, color=COLOR_H2)

    # Date and prepared by
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Date: {today}")
    _set_font(run, size=12, color=COLOR_BODY)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Prepared by: [Your Company Name]")
    _set_font(run, size=12, color=COLOR_BODY)

    # Industry tag
    industry = client_info.get("industry")
    if industry:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Industry: {industry.title()}")
        _set_font(run, size=12, bold=True, color=COLOR_H2)
        para.paragraph_format.space_before = Pt(20)

    # Page break
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Markdown Parser → DOCX
# ---------------------------------------------------------------------------

def _parse_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    """
    Parse a Markdown proposal and add content to the DOCX document.

    Handles:
      - # H1, ## H2, ### H3 → styled headings
      - - or * bullet items → bullet list
      - **bold** text → bold run
      - Regular paragraphs → body text
    """
    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # H1 heading
        if stripped.startswith("# ") and not stripped.startswith("## "):
            _add_heading(doc, stripped[2:].strip(), level=1)

        # H2 heading
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            _add_heading(doc, stripped[3:].strip(), level=2)

        # H3 heading
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), level=3)

        # Bullet item (- or *)
        elif re.match(r"^[-*]\s+", stripped):
            bullet_text = re.sub(r"^[-*]\s+", "", stripped)
            bullet_text = _strip_inline_markdown(bullet_text)
            _add_bullet(doc, bullet_text)

        # Numbered list
        elif re.match(r"^\d+\.\s+", stripped):
            item_text = re.sub(r"^\d+\.\s+", "", stripped)
            item_text = _strip_inline_markdown(item_text)
            _add_bullet(doc, item_text)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph()

        # Regular paragraph
        else:
            clean = _strip_inline_markdown(stripped)
            para = doc.add_paragraph()
            run = para.add_run(clean)
            _set_font(run, size=11, color=COLOR_BODY)
            para.paragraph_format.space_after = Pt(6)

        i += 1


def _strip_inline_markdown(text: str) -> str:
    """Remove inline markdown formatting (bold, italic, code)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)   # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)         # *italic*
    text = re.sub(r"__(.+?)__", r"\1", text)         # __bold__
    text = re.sub(r"`(.+?)`", r"\1", text)            # `code`
    return text


# ---------------------------------------------------------------------------
# Main Entry Point
# ---------------------------------------------------------------------------

def generate_docx(
    proposal_markdown: str,
    client_info: dict[str, Any],
) -> bytes:
    """
    Convert a Markdown proposal into a styled DOCX file.

    Args:
        proposal_markdown: The full proposal in Markdown format.
        client_info: ClientInfo dict for cover page data.

    Returns:
        DOCX file as bytes (not saved to disk).
    """
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover page
    _add_cover_page(doc, client_info)

    # Parse and add proposal content
    _parse_markdown_to_docx(doc, proposal_markdown)

    # Footer
    _add_footer(doc)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    docx_bytes = buffer.read()
    logger.info("DOCX generated: %d bytes", len(docx_bytes))
    return docx_bytes
